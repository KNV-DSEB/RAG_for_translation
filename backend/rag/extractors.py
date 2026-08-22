"""Trích xuất văn bản từ tệp tài liệu, giữ kèm vị trí để hiện trích dẫn.

Input:  đường dẫn tệp (.pdf, .docx, .doc, .txt, .md).
Output: `ExtractionResult` gồm danh sách khối văn bản + vị trí ("trang 3" / "đoạn 12"),
        tên bộ trích xuất đã dùng, và cờ chất lượng.

Điểm khó nhất là `.doc` (Word nhị phân cũ) — 2/3 tài liệu kiểm thử thật thuộc loại này.
Xử lý theo chuỗi 3 lớp, dừng ở lớp đầu tiên thành công:
    1. Word COM (máy có Word 16) — tin cậy nhất
    2. Trích thô từ cấu trúc OLE — lưới an toàn, gắn cờ chất lượng thấp
    3. Thất bại cả hai → báo lỗi cụ thể, KHÔNG nạp tài liệu rỗng rồi trả lời sai
"""

from __future__ import annotations

import re
import unicodedata
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

from backend.config import settings

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".doc", ".txt", ".md"})

# Word COM có thể treo nếu gặp hộp thoại ẩn — luôn đặt hạn giờ.
WORD_COM_TIMEOUT_SEC = 60

# Ký tự coi là "văn bản hợp lệ": ASCII in được + Latin mở rộng + dải dấu tiếng Việt.
_PLAUSIBLE_RE = re.compile(r"[\x20-\x7EÀ-ɏḀ-ỿ\s]")
# Ngưỡng tỷ lệ ký tự hợp lệ để giữ lại một dòng khi trích thô từ OLE.
_PLAUSIBLE_RATIO = 0.85


@dataclass
class ExtractedBlock:
    """Một khối văn bản kèm vị trí gốc, dùng để dựng trích dẫn bấm mở được."""

    text: str
    locator: str


@dataclass
class ExtractionResult:
    blocks: list[ExtractedBlock]
    extractor: str
    quality: Literal["ok", "low"] = "ok"
    warnings: list[str] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks)

    @property
    def n_chars(self) -> int:
        return sum(len(block.text) for block in self.blocks)


# ============================== Lỗi có nghĩa nghiệp vụ ==============================


class ExtractionError(RuntimeError):
    """Lỗi trích xuất, thông báo đã bằng tiếng Việt và nói rõ làm gì tiếp."""


class UnsupportedFormat(ExtractionError):
    pass


class EncryptedFile(ExtractionError):
    pass


class ScannedPdf(ExtractionError):
    pass


class EmptyDocument(ExtractionError):
    pass


# ============================== Tiện ích ==============================


def _normalize(text: str) -> str:
    """Chuẩn hoá Unicode về NFC để chữ tiếng Việt có dấu so khớp và hiển thị đúng."""
    return unicodedata.normalize("NFC", text).replace("\r\n", "\n").replace("\r", "\n")


def _is_plausible_text(line: str) -> bool:
    """Dòng có phải văn bản thật không, hay là ký tự rác từ cấu trúc nhị phân."""
    stripped = line.strip()
    if len(stripped) < 12:
        return False
    good = len(_PLAUSIBLE_RE.findall(stripped))
    return (good / len(stripped)) >= _PLAUSIBLE_RATIO


def _blocks_from_paragraphs(paragraphs: list[str]) -> list[ExtractedBlock]:
    """Gom các đoạn không rỗng thành khối, đánh số vị trí theo thứ tự đoạn."""
    blocks: list[ExtractedBlock] = []
    index = 0
    for para in paragraphs:
        text = _normalize(para).strip()
        if not text:
            continue
        index += 1
        blocks.append(ExtractedBlock(text=text, locator=f"đoạn {index}"))
    return blocks


# ============================== PDF ==============================


def _extract_pdf(path: Path) -> ExtractionResult:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
    except PdfReadError as exc:
        raise ExtractionError(
            f"Không đọc được tệp PDF '{path.name}' — tệp có thể bị hỏng. "
            "Thử mở bằng trình đọc PDF và lưu lại, rồi nạp lần nữa."
        ) from exc

    if reader.is_encrypted:
        # Thử mật khẩu rỗng: nhiều PDF chỉ khoá quyền in/sửa chứ không khoá đọc.
        try:
            if reader.decrypt("") == 0:
                raise EncryptedFile(
                    f"Tệp '{path.name}' được bảo vệ bằng mật khẩu. "
                    "Bỏ mật khẩu rồi nạp lại."
                )
        except EncryptedFile:
            raise
        except Exception as exc:
            raise EncryptedFile(
                f"Tệp '{path.name}' được bảo vệ bằng mật khẩu. Bỏ mật khẩu rồi nạp lại."
            ) from exc

    blocks: list[ExtractedBlock] = []
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = _normalize(page.extract_text() or "").strip()
        except Exception:
            text = ""
        if text:
            blocks.append(ExtractedBlock(text=text, locator=f"trang {page_no}"))

    total = sum(len(b.text) for b in blocks)
    if total < settings.scanned_pdf_char_threshold:
        # Không nạp tài liệu rỗng rồi âm thầm trả lời sai (spec edge case §4.1).
        raise ScannedPdf(
            f"Tệp '{path.name}' hầu như không có lớp chữ — nhiều khả năng là bản scan/ảnh chụp. "
            f"Chỉ đọc được {total} ký tự trên {len(reader.pages)} trang. "
            "MVP này chưa làm OCR: hãy dùng bản có chữ, hoặc chuyển sang tệp Word."
        )

    return ExtractionResult(blocks=blocks, extractor="pypdf")


# ============================== DOCX ==============================


def _extract_docx(path: Path) -> ExtractionResult:
    import docx  # python-docx

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(
            f"Không đọc được tệp '{path.name}'. Nếu đây là tệp .doc cũ được đổi đuôi thành "
            ".docx thì hãy đổi lại đúng đuôi .doc rồi nạp."
        ) from exc

    parts: list[str] = [para.text for para in document.paragraphs]

    # Bảng biểu: ghép từng hàng thành một dòng, giữ được số liệu trong bảng.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    blocks = _blocks_from_paragraphs(parts)
    if not blocks:
        raise EmptyDocument(f"Tệp '{path.name}' không có nội dung chữ để lập chỉ mục.")
    return ExtractionResult(blocks=blocks, extractor="python_docx")


# ============================== DOC (Word cũ) ==============================


def _doc_via_word_com(path: Path) -> list[str]:
    """Lớp 1: dùng Word COM chuyển .doc sang .docx trong thư mục tạm rồi đọc.

    Chạy trong luồng riêng có CoInitialize và LUÔN Quit() trong finally để không
    bỏ lại tiến trình Word treo. Dùng DispatchEx (KHÔNG dùng Dispatch) để tạo một
    instance Word riêng — nếu bám vào Word người dùng đang mở thì Quit() sẽ đóng
    luôn tài liệu họ đang làm dở.
    """
    import tempfile

    import pythoncom  # type: ignore[import-not-found]
    import win32com.client  # type: ignore[import-not-found]

    WD_FORMAT_DOCX = 16

    pythoncom.CoInitialize()
    word = None
    tmp_dir = tempfile.mkdtemp(prefix="docconv_")
    tmp_docx = Path(tmp_dir) / (path.stem + ".docx")
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        try:
            document.SaveAs2(str(tmp_docx), FileFormat=WD_FORMAT_DOCX)
        finally:
            document.Close(SaveChanges=False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    import docx

    converted = docx.Document(str(tmp_docx))
    parts = [p.text for p in converted.paragraphs]
    for table in converted.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    try:
        tmp_docx.unlink(missing_ok=True)
        Path(tmp_dir).rmdir()
    except OSError:
        pass

    return parts


def _doc_via_ole(path: Path) -> list[str]:
    """Lớp 2: trích thô từ cấu trúc OLE.

    Văn bản trong .doc thường nằm ở stream `WordDocument`, mã hoá UTF-16LE xen lẫn
    dữ liệu nhị phân. Giải mã rồi lọc theo tỷ lệ ký tự hợp lệ để bỏ phần rác.
    Cách này đã thử nghiệm thành công trên đúng 3 tệp kiểm thử của dự án.
    """
    import olefile

    raw = b""
    if olefile.isOleFile(str(path)):
        with olefile.OleFileIO(str(path)) as ole:
            for stream in ("WordDocument", "1Table", "0Table"):
                if ole.exists(stream):
                    raw += ole.openstream(stream).read()
    if not raw:
        raw = path.read_bytes()

    decoded = raw.decode("utf-16-le", errors="ignore")
    # Ranh giới đoạn trong .doc hay là CR; tách theo cả CR/LF và ký tự điều khiển.
    candidates = re.split(r"[\r\n\x00-\x08\x0b\x0c\x0e-\x1f]+", decoded)

    parts: list[str] = []
    for candidate in candidates:
        cleaned = candidate.strip()
        if _is_plausible_text(cleaned):
            parts.append(cleaned)
    return parts


def _extract_doc(path: Path) -> ExtractionResult:
    warnings: list[str] = []

    # --- Lớp 1: Word COM ---
    try:
        with ThreadPoolExecutor(max_workers=1) as pool:
            parts = pool.submit(_doc_via_word_com, path).result(timeout=WORD_COM_TIMEOUT_SEC)
        blocks = _blocks_from_paragraphs(parts)
        if blocks:
            return ExtractionResult(blocks=blocks, extractor="word_com", quality="ok")
        warnings.append("Word đọc được tệp nhưng không thấy nội dung chữ nào.")
    except Exception as exc:
        warnings.append(f"Không dùng được Word để đọc tệp ({type(exc).__name__}).")

    # --- Lớp 2: trích thô từ OLE ---
    try:
        parts = _doc_via_ole(path)
        blocks = _blocks_from_paragraphs(parts)
        if blocks:
            warnings.append(
                "Đã đọc bằng cách trích thô vì Word không dùng được — "
                "chất lượng có thể thấp, nên bấm Xem trước để kiểm tra nội dung."
            )
            return ExtractionResult(
                blocks=blocks, extractor="ole_fallback", quality="low", warnings=warnings
            )
    except Exception as exc:
        warnings.append(f"Trích thô cũng thất bại ({type(exc).__name__}).")

    # --- Lớp 3: chịu thua, báo rõ ---
    raise ExtractionError(
        f"Không đọc được nội dung tệp '{path.name}'. "
        + " ".join(warnings)
        + " Cách xử lý: mở tệp bằng Word rồi lưu lại dạng .docx và nạp bản đó."
    )


# ============================== TXT / MD ==============================


def _extract_plain(path: Path) -> ExtractionResult:
    text = ""
    for encoding in ("utf-8", "utf-8-sig", "cp1258", "latin-1"):
        try:
            text = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if not text.strip():
        raise EmptyDocument(f"Tệp '{path.name}' rỗng hoặc không đọc được nội dung.")

    # Tách theo dòng trống để giữ ranh giới đoạn tự nhiên.
    paragraphs = re.split(r"\n\s*\n", _normalize(text))
    blocks = _blocks_from_paragraphs(paragraphs)
    if not blocks:
        raise EmptyDocument(f"Tệp '{path.name}' không có nội dung chữ để lập chỉ mục.")
    return ExtractionResult(blocks=blocks, extractor="plain")


# ============================== Điểm vào ==============================


def extract(path: Path | str) -> ExtractionResult:
    """Trích xuất văn bản từ một tệp. Ném `ExtractionError` với lý do cụ thể nếu hỏng."""
    path = Path(path)
    if not path.exists():
        raise ExtractionError(f"Không tìm thấy tệp '{path.name}'.")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFormat(
            f"Chưa hỗ trợ định dạng '{ext}'. Các định dạng nhận được: {supported}."
        )

    if path.stat().st_size == 0:
        raise EmptyDocument(f"Tệp '{path.name}' rỗng (0 byte).")

    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".doc":
        return _extract_doc(path)
    return _extract_plain(path)
